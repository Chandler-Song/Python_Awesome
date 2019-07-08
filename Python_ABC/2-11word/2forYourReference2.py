import docx
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE

document = docx.Document('test.docx')

table = document.tables[0]

for row, obj_row in enumerate(table.rows):
    for col, cell in enumerate(obj_row.cells):
        cell.text = '{}, {}'.format(row, col)

table.columns[4].width = docx.shared.Cm(4)

# styles = document.styles
# table_style = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
# for style in table_style:
# 	print(style.name)
table.style = 'Table Grid'

row = table.rows[1]
row.cells[0].text = 'elephant'
row.cells[1].text = 'shark'

row_count = len(table.rows)
col_count = len(table.columns)
print('row = {}, column = {}'.format(row_count, col_count))

# row = table.add_row()
document.add_page_break()

# Adding a  picture
document.add_picture('classAndObject.png', width = Cm(5))

document.save('test.docx')





