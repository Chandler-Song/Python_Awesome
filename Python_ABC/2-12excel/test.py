import openpyxl
import docx

def rowDataSet(ws, row):

	rowSet = set()	# 空集合

	for cell in ws[row]:
		if cell.value not in [None, '']: # 单元格不为空，把内容加入集合
			rowSet.add(cell.value)

	return rowSet	# 返回集合

# 输出第一个表格
def exportDataToWordTable1(ws, excelRow, doc, analysis):

	# 画word里面的表格（两行七列）
	docRows = 2
	docCols = 7
	table = doc.add_table(rows = docRows, cols = docCols, style = 'Table Grid')

	# excel数据填入word表格1里
	for row in range(docRows):
		for col in range(docCols):
			table.rows[row].cells[col].text = str(ws.cell(row=excelRow, column= col+1).value)

		excelRow += 1

	table.rows[0].cells[1].text = '票房'
	table.rows[0].cells[2].text = '天数'
	table.rows[0].cells[3].text = '票房／天数'
	table.rows[1].cells[3].text = '{:.3f}'.format(float(table.rows[1].cells[1].text)/float(table.rows[1].cells[2].text))
	table.rows[0].cells[1].text = ''

	# 数据分析，保存结果
	if float(table.rows[1].cells[4].text) > 0.85 and float(table.rows[1].cells[5].text) > 0.85:
		analysis['direction'] = '南'
	else:
		analysis['direction'] = '北'

	if float(table.rows[1].cells[6].text) < 0.2:
		analysis['feeling'] = '凉'
	else:
		analysis['feeling'] = '暖'

	excelRow += 1 # 跳过factor_loading那行
	vSet = set()

	# facotr_loading到alpha这一列的数据放到vSet集合里
	while str(ws.cell(row=excelRow, column=2).value) != 'alpha':
		v = ws.cell(row=excelRow, column=2).value
		vSet.add(abs(v))
		excelRow += 1

	if min(vSet) > 0.3:
		analysis['season'] = '夏'
		analysis['duration'] = '长'
	else:
		analysis['season'] = '冬'
		analysis['duration'] = '短'

	# 一二表格间加一空行
	doc.add_paragraph('\n')

	return excelRow, analysis

# 输出第二个表格
def exportDataToWordTable2(ws, excelRow, doc, analysis):

	docCols = 2
	# 一行一行往word里面添加表格2的数据
	table = doc.add_table(rows=1, cols=docCols, style='Table Grid')
	wordRow = 0  # word表格里的行指针

	# excelRow从alpha这一行开始，到X这一行结束。从excel一行一行去数据，填入word表格
	while ws.cell(row=excelRow, column=1).value != 'X':
		for docCol in range(docCols):
			table.rows[wordRow].cells[docCol].text = str(ws.cell(row=excelRow, column=docCol+1).value)
		wordRow += 1
		excelRow += 1
		table.add_row()

	# 记录alpha的值，并进行分析
	analysis['alpha'] = float(table.rows[1].cells[1].text)

	if analysis['alpha'] > 0.8:
		analysis['climate'] = '炎热'
	else:
		analysis['climate'] = '严寒'

	doc.add_paragraph('\n')

	return excelRow, analysis

# 输出第三个表格
def exportDataToWordTable3(ws, excelRow, doc):

	# 首先确定行列
	docCols = len(rowDataSet(ws, excelRow))

	# word里画出表格
	table = doc.add_table(rows=docCols, cols=docCols, style='Table Grid')

	# 提取excel数据填入word表格3
	for docRow in range(docCols):
		for docCol in range(docCols):
			table.rows[docRow].cells[docCol].text = str(ws.cell(row=excelRow, column=docCol+1).value)
		excelRow += 1

	doc.add_paragraph('\n')

	return excelRow

# 文字输出（对表格数据的分析）
def textOutput(doc, analysis):

	outputStr = '''人皆苦{}，我爱{}日{}。 \n熏风自{}来，殿阁生微{}。 \n
		alpha = {}'''.format(analysis['climate'], analysis['season'],
							 analysis['duration'], analysis['direction'],
							 analysis['feeling'], analysis['alpha'])

	p = doc.add_paragraph(outputStr)
	p.add_run('\n')

	doc.add_paragraph('-*-'*33)
	doc.add_page_break()

wb = openpyxl.load_workbook('data.xlsx')
ws = wb.get_sheet_by_name('Sheet 1')

doc = docx.Document()

dataAnalysis = {'direction':'',
				'feeling':'',
				'season':'',
				'duration':'',
				'climate':'',
				'alpha':0}
# 初始化excel行指针
row = 1

while row < ws.max_row:

	if {'TLI', 'CFI', 'RMSEA'} < rowDataSet(ws, row):

		# 输出第一张表，记录数据分析结果
		row, dataAnalysis = exportDataToWordTable1(ws, row, doc, dataAnalysis)

		# 输出第二张表，记录数据分析结果
		row, dataAnalysis = exportDataToWordTable2(ws, row, doc, dataAnalysis)

		# 输出第三张表，记录数据分析结果
		row = exportDataToWordTable3(ws, row, doc)

		# 输出数据分析结果
		textOutput(doc, dataAnalysis)

	row += 1

doc.save('analysisResult.docx')





